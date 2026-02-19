from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# Helper functions
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return h

def add_para(text, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_rich_para(parts, space_after=6):
    """parts is list of (text, bold, italic) tuples"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    for text, bold, italic in parts:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
    return p

def add_bullet(text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()  # spacing after table
    return table

def add_separator():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('\u2500' * 60)
    run.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
    run.font.size = Pt(8)


# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('AI Manager Skills Platform')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Commercial Viability Review')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

version = doc.add_paragraph()
version.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = version.add_run('Revision 2 \u2014 Post-Hardening Update')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x44, 0x88, 0xcc)
run.bold = True

doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run(datetime.date.today().strftime('%B %d, %Y'))
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_paragraph()
doc.add_paragraph()

conf = doc.add_paragraph()
conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = conf.add_run('Updated review reflecting platform-hardening changes: schema validation, '
                    'accessibility improvements, onboarding, and expanded assessment criteria')
run.italic = True
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_page_break()


# ============================================================
# WHAT CHANGED SINCE V1
# ============================================================
add_heading('What Changed Since the Initial Review', level=1)

add_para(
    'The initial review (February 15, 2026) identified gaps across security, accessibility, '
    'onboarding, and assessment rigor. A platform-hardening sprint addressed several of these. '
    'This section documents what was fixed and what remains open.'
)

add_table(
    ['Issue from v1 Review', 'Status', 'What Changed'],
    [
        ['get_optional_user swallows all exceptions',
         'FIXED',
         'Now catches specific JWTError; logs unexpected errors with context'],
        ['No file size/type validation on uploads',
         'FIXED',
         '5 MB max file size enforced; only .json and .txt extensions allowed'],
        ['No focus trapping on mobile sidebar',
         'FIXED',
         'Full focus trap with Tab/Shift+Tab cycling, auto-focus on open, Escape to close'],
        ['No onboarding or guided path',
         'FIXED',
         'GettingStartedOverlay component added with 3-step walkthrough for new users'],
        ['No "what to do next" recommendations',
         'FIXED',
         'useRecommendedLesson hook suggests next incomplete lesson based on optimal learning order'],
        ['No lang attribute on HTML',
         'FIXED',
         'lang="en" present on <html> element'],
        ['Schema inputs lack validation',
         'FIXED',
         'Field(max_length) validators added across all lesson schemas (9 schema files)'],
        ['Assessment criteria too few per lesson',
         'IMPROVED',
         'Expanded from 3 to 4 criteria per lesson; stricter completion checks in backend'],
        ['Tokens in localStorage (XSS vulnerable)',
         'OPEN',
         'Still uses localStorage for JWT tokens'],
        ['No SSO/SAML integration',
         'OPEN',
         'JWT-only authentication remains'],
        ['No password reset flow',
         'OPEN',
         'Only authenticated change-password exists'],
        ['14 window.confirm() dialogs',
         'OPEN',
         'ConfirmDialog component exists but only used in admin; lessons still use confirm()'],
        ['No payment/subscription system',
         'OPEN',
         'Zero monetization infrastructure'],
        ['No certification at completion',
         'OPEN',
         'No certificate generation or competency proof'],
        ['High-contrast theme not enabled',
         'OPEN',
         'CSS defined but no UI toggle to activate it'],
    ]
)

doc.add_page_break()


# ============================================================
# EXECUTIVE SUMMARY (UPDATED)
# ============================================================
add_heading('Executive Summary', level=1)

add_rich_para([
    ('Could this be sold as an educational tool? ', False, False),
    ('Yes \u2014 and it is closer now than it was a week ago.', True, False),
], space_after=8)

add_para(
    'The core strengths identified in the initial review remain: the six managerial concepts are '
    'well-defined and teachable, the interactive tools provide genuine hands-on practice, and the '
    'tech stack is production-quality. The platform-hardening sprint closed several security and UX '
    'gaps, most notably input validation across all schemas, accessibility improvements on the mobile '
    'sidebar, and a proper onboarding flow for new users.'
)

add_para(
    'The remaining gaps are concentrated in the business layer: no payment system, no certification, '
    'no SSO, and assessment thresholds that remain trivially low despite being expanded. The path to '
    'Tier 1 (self-serve individual product) is now shorter \u2014 estimated at 2\u20133 weeks rather '
    'than the original 2\u20134 weeks.'
)

doc.add_page_break()


# ============================================================
# WHAT'S GENUINELY GOOD (retained from v1, minor updates)
# ============================================================
add_heading("What's Genuinely Good", level=1)

add_heading('The Curriculum Is the Real Asset', level=2)

add_para(
    'The six managerial concepts (Context Assembly, Quality Judgment, Task Decomposition, '
    'Iterative Refinement, Workflow Integration, Frontier Recognition) are well-defined, '
    'teachable, and address real pain points. A corporate training buyer would find these '
    'credible \u2014 they are not buzzwords.'
)

add_para(
    'Every lesson follows a tight learn-practice-track cycle. Users build actual deliverables '
    '(templates, trust matrices, checklists, delegation plans, reference cards), not just answer '
    'quizzes. The Learn tabs use concrete workplace scenarios across marketing, HR, finance, IT, '
    'and operations \u2014 not just tech examples.'
)

add_table(
    ['Module', 'Lessons', 'Primary Concepts'],
    [
        ['Foundation', '1\u20133', 'Context Assembly'],
        ['Documentation & Trust', '4\u20136', 'Context Assembly, Quality Judgment'],
        ['Workflow', '7\u201310', 'Task Decomposition, Iterative Refinement, Workflow Integration'],
        ['Advanced', '11\u201312', 'Frontier Recognition, Integration'],
    ]
)

add_heading('Pedagogical Design Is Solid', level=2)

add_bullet('Problem/Skill framing: ', 'Immediate relevance \u2014 ', 0)
add_para(
    'Each lesson opens with "The Problem" and "The Skill," immediately answering '
    '"why should I care?" before diving into content.'
)

add_bullet('Behavioral assessment: ', 'Measures doing, not knowing \u2014 ', 0)
add_para(
    'Self-assessment criteria are behavioral ("analyzed at least 1 conversation," '
    '"identified at least 1 context pattern") rather than knowledge-based. '
    'Now expanded to 4 criteria per lesson with stricter completion logic.'
)

add_bullet('Cross-lesson integration: ', 'Lessons reference each other \u2014 ', 0)
add_para(
    'Lesson 3 pulls data from Lesson 1. Lesson 8 imports from Lesson 7. Lesson 12 aggregates '
    'all lessons. ConnectionCallout components create a narrative thread without forcing '
    'sequential unlocking.'
)

add_bullet('Before/after comparisons: ', 'Show, don\'t tell \u2014 ', 0)
add_para(
    'Every lesson includes side-by-side examples of vague vs. specific approaches. '
    'The differences are visceral and memorable.'
)

add_heading('Technology Stack Is Production-Quality', level=2)

add_table(
    ['Layer', 'Technology', 'Notes'],
    [
        ['Backend', 'FastAPI (Python, async)', 'JWT auth, rate limiting, circuit breaker, schema validation'],
        ['AI Integration', 'Claude API (Anthropic)', 'Real LLM calls with retry logic and graceful degradation'],
        ['Database', 'PostgreSQL + SQLAlchemy 2.0', '29 models, UUID primary keys, per-user data isolation, cascade deletes'],
        ['Frontend', 'React 18 + Vite', 'Custom CSS (4,188 lines), dark/light themes, responsive, a11y improvements'],
        ['Deployment', 'Docker Compose', 'Frontend (nginx), backend, postgres, admin panel'],
    ]
)

add_para(
    'The AI integration includes a circuit breaker pattern (opens after 5 failures, half-open '
    'after 30s cooldown), exponential backoff retries (1s, 2s, 4s), and structured error handling. '
    'All Pydantic schemas now enforce max_length constraints on string fields.'
)

add_heading('New: Onboarding and Guided Path', level=2)

add_para(
    'A GettingStartedOverlay now greets first-time users with a 3-step introduction: '
    '"Pick a Lesson," "Practice the Skill," and "Complete the Checklist." After dismissal, '
    'a recommended-lesson banner on the Dashboard suggests the next incomplete lesson based on '
    'an optimal learning order (1, 3, 5, 7, 2, 4, 6, 8, 9, 10, 11, 12). Both features use '
    'localStorage to persist state.'
)

doc.add_page_break()


# ============================================================
# WHAT'S STILL MISSING
# ============================================================
add_heading("What's Still Missing for a Sellable Product", level=1)

add_heading('1. Assessment Rigor Remains Too Low', level=2)

add_para(
    'Assessment was expanded from 3 to 4 criteria per lesson, and the backend now enforces '
    'stricter completion logic (e.g., Lesson 1 requires edited notes, Lesson 2 requires entries '
    'spanning 3+ categories). However, the fundamental issue remains: criteria are existence-based, '
    'not quality-based. Creating 1 template of any quality passes Lesson 3.'
)

add_bullet('4 criteria per lesson, 48 total \u2014 but all are binary (did/didn\'t)')
add_bullet('Only Lesson 2 (Feedback Analyzer) applies quality scoring to user work')
add_bullet('No rubric, no peer review, no minimum depth checks')
add_bullet('No final integrated competency assessment')
add_bullet('A corporate buyer paying per-seat still cannot measure learning outcomes')

add_heading('2. No Certification or Competency Proof', level=2)

add_para(
    'Completing all 12 lessons generates a progress percentage. There is no certificate, badge, '
    'downloadable credential, or verifiable proof of completion. This is a blocker for enterprise '
    'L&D buyers who need documented training records.'
)

add_heading('3. Organizational Features Are Half-Built', level=2)

add_para(
    'The admin backend has cohort management API endpoints (create, update, add/remove members, '
    'basic stats). However, there is no admin UI frontend. No facilitator view, no team progress '
    'dashboard, no group exercises, no peer review. The infrastructure exists but is unexposed.'
)

add_bullet('Admin API: cohort CRUD, experiment management, user listing \u2014 all functional')
add_bullet('Admin UI: does not exist \u2014 no frontend consumes the admin endpoints')
add_bullet('No manager dashboard, no discussion forums, no facilitation guides')

add_heading('4. Remaining Security Gaps', level=2)

add_table(
    ['Issue', 'Severity', 'Fix Effort', 'Status'],
    [
        ['Tokens in localStorage (XSS vulnerable)', 'High', 'Medium \u2014 switch to HttpOnly cookies', 'OPEN'],
        ['Rate limiting on only 2\u20133 endpoints', 'Medium', 'Low \u2014 extend to all AI endpoints', 'OPEN'],
        ['No SSO/SAML integration', 'High (enterprise)', 'High \u2014 requires auth refactor', 'OPEN'],
        ['No password reset flow', 'Medium', 'Low \u2014 add email flow', 'OPEN'],
        ['14 window.confirm() dialogs', 'Low', 'Medium \u2014 replace with accessible modals', 'OPEN'],
        ['High-contrast theme not wired up', 'Low', 'Low \u2014 add toggle in ThemeContext', 'OPEN'],
        ['File size/type validation on uploads', 'Was Medium', 'N/A', 'FIXED'],
        ['get_optional_user exception handling', 'Was Low', 'N/A', 'FIXED'],
        ['Schema string validation', 'Was Medium', 'N/A', 'FIXED'],
    ]
)

add_heading('5. Content Gaps', level=2)

add_bullet('No "apply this to your real work this week" structured exercises')
add_bullet('No advanced/stretch path for fast learners')
add_bullet('No industry customization (everyone sees same examples regardless of role)')
add_bullet('No frontier case library from real encounters')

add_para(
    'Note: The v1 review claimed Lesson 11 (Frontier Mapper) "feels thinner than the others." '
    'Code audit shows this is inaccurate \u2014 Lesson 11 has a 682-line backend (3rd largest), '
    'extensive Learn tab content, and full CRUD with statistics. It is comparable to Lessons 3 and 5.',
    italic=True
)

add_heading('6. No Monetization Infrastructure', level=2)

add_bullet('No payment/subscription system (Stripe, etc.)')
add_bullet('No license management or seat counting')
add_bullet('No usage metering')
add_bullet('No guest-to-paid conversion funnel')

doc.add_page_break()


# ============================================================
# DETAILED TECHNICAL FINDINGS (UPDATED)
# ============================================================
add_heading('Detailed Technical Findings', level=1)

add_heading('Backend Architecture', level=2)

add_para(
    'The backend is well-structured with a modular design: each lesson has its own directory '
    'with schemas, routes, and analyzers. All endpoints have real business logic \u2014 none are stubs.'
)

add_rich_para([
    ('AI Integration: ', True, False),
    ('Real Claude API calls (claude-3-haiku by default, configurable). Integrated in Lessons 1, 3, 5, and 12. '
     'Circuit breaker tracks failures in 60-second windows, opens after 5 failures, allows test request '
     'after 30s cooldown. Retries up to 3 times with exponential backoff (1s, 2s, 4s). Auth errors never retried.', False, False),
])

add_rich_para([
    ('Authentication: ', True, False),
    ('Production-ready JWT with access tokens (15-min) and refresh tokens (7-day). Passwords hashed with '
     'bcrypt (12 rounds). Rate limiting on register/login (5/minute). Token type validation prevents '
     'refresh-as-access attacks. Refresh token revocation on logout. Exception handling in get_optional_user '
     'now catches specific JWTError and logs unexpected errors.', False, False),
])

add_rich_para([
    ('Input Validation: ', True, False),
    ('All Pydantic schemas across 9 lesson modules now enforce Field(max_length) constraints on string fields. '
     'File uploads validated for size (5 MB max) and type (.json, .txt only). This prevents oversized payloads '
     'and unexpected file types from reaching the database.', False, False),
])

add_rich_para([
    ('Database: ', True, False),
    ('PostgreSQL with SQLAlchemy 2.0 async. UUID primary keys. Per-user data isolation via user_id foreign keys '
     'with CASCADE deletes. 29 models covering all 12 lessons plus analytics, admin cohorts, and A/B experiments.', False, False),
])

add_rich_para([
    ('Test Coverage: ', True, False),
    ('18 test files totaling approximately 3,150 lines. Tests exist for auth (including change-password), '
     'error paths, and all 12 lessons. Uses pytest + pytest-asyncio. Estimated 30\u201340% coverage. '
     'Still needs expansion for production confidence, particularly integration tests.', False, False),
])

add_heading('Frontend User Experience', level=2)

add_para(
    'React 18 with modern hooks, React Router v6, Vite 5 build tool. Zero external UI libraries \u2014 '
    'all custom CSS (4,188 lines, well-organized with CSS custom properties for theming).'
)

add_rich_para([
    ('Themes: ', True, False),
    ('Light and dark themes via CSS custom properties (~70 unique variables, ~210 total declarations across '
     '3 theme blocks). High-contrast theme fully defined in CSS but not yet wired to a UI toggle.', False, False),
])

add_rich_para([
    ('Responsive Design: ', True, False),
    ('Three breakpoints (768px, 1024px, 480px). Mobile hamburger menu with slide-in sidebar. '
     'Grid layouts reflow naturally. Touch targets meet WCAG minimum (44px). No horizontal scroll issues.', False, False),
])

add_rich_para([
    ('Accessibility: ', True, False),
    ('Skip link, lang="en" on html, aria-expanded on toggles, aria-label on icon buttons, '
     'focus-visible states, reduced-motion support. Mobile sidebar now has full focus trapping '
     '(Tab/Shift+Tab cycling, auto-focus on open). Remaining gaps: 14 window.confirm() dialogs '
     'in lesson pages (ConfirmDialog component exists but is only used in admin), and high-contrast '
     'theme is not accessible from the UI. All form inputs use proper <label> elements.', False, False),
])

add_rich_para([
    ('Onboarding: ', True, False),
    ('New GettingStartedOverlay provides a 3-step walkthrough for first-time users. '
     'useRecommendedLesson hook displays a "Recommended Next" banner on the Dashboard, '
     'guiding users through an optimal lesson sequence (1, 3, 5, 7, 2, 4, 6, 8, 9, 10, 11, 12).', False, False),
])

add_heading('Lesson Experience (What Users Actually Do)', level=2)

add_table(
    ['Lesson', 'Title', 'What Users Build', 'AI-Powered?'],
    [
        ['1', 'Context Tracker', 'Analyzed conversations with coaching feedback', 'Yes \u2014 full analysis'],
        ['2', 'Feedback Analyzer', 'Scored feedback entries with rewrite suggestions', 'Rule-based'],
        ['3', 'Template Builder', 'Reusable prompt templates with test results', 'Yes \u2014 template testing'],
        ['4', 'Context Docs', 'Living project documents with AI session prompts', 'Prompt generation'],
        ['5', 'Trust Matrix', 'Personal trust ratings with prediction tracking', 'Yes \u2014 calibration insights'],
        ['6', 'Verification Tools', 'Reusable verification checklists', 'CRUD only'],
        ['7', 'Task Decomposer', 'Project breakdowns with AI/human categorization', 'CRUD + analytics'],
        ['8', 'Delegation Tracker', 'Structured delegations with task workflows', 'CRUD + templates'],
        ['9', 'Iteration Passes', 'Multi-pass refinement records with feedback quality', 'Feedback scoring'],
        ['10', 'Status Reporter', 'Workflow templates with execution tracking', 'Prompt generation'],
        ['11', 'Frontier Mapper', 'AI reliability zone maps with encounter logs', 'CRUD + statistics'],
        ['12', 'Reference Card', 'Personal quick-reference aggregating all lessons', 'Yes \u2014 card generation'],
    ]
)

add_heading('Assessment Criteria (Updated)', level=2)

add_para(
    'Each lesson now has exactly 4 completion criteria, up from 3 in the initial build. '
    'The backend progress endpoint enforces these with stricter completion logic. Examples:'
)

add_table(
    ['Lesson', 'Criteria (4 per lesson)', 'New in v2'],
    [
        ['1 \u2014 Context Tracker', 'Analyzed 1+ conversation, identified pattern, reviewed coaching, edited notes', 'Edited notes'],
        ['2 \u2014 Feedback Analyzer', 'Analyzed 3+ entries, rewrote 1+ vague, reviewed patterns, 3+ categories', '3+ categories'],
        ['3 \u2014 Template Builder', 'Created 1+ template, tested 1+, used variables, rated 1+ test', 'Rated test'],
        ['5 \u2014 Trust Matrix', 'Created 3+ output types, tracked 1+ prediction, reviewed calibration, 2+ domains', '2+ domains'],
        ['7 \u2014 Task Decomposer', 'Created 1+ decomposition, 5+ tasks, categorized AI/human, dependency chain', 'Dependency chain'],
        ['12 \u2014 Reference Card', 'Generated card, filled 3+ sections, exported, completed challenge', 'Challenge completion'],
    ]
)

add_para(
    'While these are stricter than the initial 3-criterion setup, they remain existence-based. '
    'A single low-quality artifact still satisfies most criteria. Quality rubrics would require '
    'either AI-powered evaluation or peer review \u2014 neither is implemented.',
    italic=True
)

doc.add_page_break()


# ============================================================
# COMPETITIVE POSITIONING
# ============================================================
add_heading('Competitive Positioning', level=1)

add_table(
    ['Factor', 'This Project', 'Typical Competitors'],
    [
        ['Content depth', 'Strong \u2014 12 hands-on lessons', 'Often shallow webinars or slide decks'],
        ['Interactivity', 'Real tools with AI integration', 'Usually just videos + quizzes'],
        ['Personalization', 'Builds personal artifacts', 'Generic one-size-fits-all'],
        ['Onboarding', 'Getting started overlay + recommended path', 'Varies \u2014 some good, some absent'],
        ['Enterprise readiness', 'Weak \u2014 no SSO, no cohort UI, no certs', 'Strong in established LMS platforms'],
        ['Price point fit', '$20\u201350 self-serve, $200\u2013500/seat enterprise', 'Varies widely'],
    ]
)

add_rich_para([
    ('The differentiator is real. ', True, False),
    ('Most "AI training" is either deeply technical (prompt engineering for developers) or superficially '
     'motivational ("AI will change everything!"). This occupies a genuine gap: ', False, False),
    ('practical skill-building for non-technical knowledge workers.', True, False),
    (' That is a real market.', False, False),
])

add_heading('Target Buyer Fit', level=2)

add_rich_para([('Strong fit:', True, False)])
add_bullet('Knowledge workers wanting to skill up individually on AI collaboration')
add_bullet('Managers wanting personal AI skills before rolling out to teams')
add_bullet('Career development for non-technical roles moving into AI-adjacent work')
add_bullet('Self-paced learner populations (remote-first, asynchronous)')

add_rich_para([('Weaker fit:', True, False)])
add_bullet('Organizations needing to upskill entire teams at once (cohort API exists but no UI)')
add_bullet('Regulated industries needing documented competency verification')
add_bullet('Companies wanting trainer-led cohorts with peer learning')
add_bullet('Groups wanting industry-specific content only')

doc.add_page_break()


# ============================================================
# WHAT IT WOULD TAKE TO SELL (UPDATED)
# ============================================================
add_heading('What It Would Take to Sell', level=1)

add_heading('Tier 1 \u2014 Self-Serve Individual Product ($20\u201350/user)', level=2)

add_bullet('Add Stripe integration and paywall')
add_bullet('Fix security basics (HttpOnly cookies for tokens)', bold_prefix='Reduced scope: ')
add_bullet('Add password reset via email')
add_bullet('Replace 14 window.confirm() calls with accessible modal dialogs')
add_bullet('Wire up high-contrast theme toggle')
add_bullet('Polish guest-to-paid conversion funnel')
add_para(
    'Estimated effort: 2\u20133 weeks (reduced from 2\u20134 \u2014 onboarding and upload validation already done)',
    italic=True
)

add_heading('Tier 2 \u2014 Team/SMB Product ($200\u2013500/seat)', level=2)

add_bullet('Everything in Tier 1')
add_bullet('Build admin UI frontend for existing cohort API endpoints')
add_bullet('Completion certificates (PDF generation)')
add_bullet('Raise assessment thresholds and add quality rubrics (AI-scored or peer review)')
add_bullet('Add "apply to your work" structured exercises')
add_bullet('SSO/SAML for enterprise authentication')
add_para('Estimated effort: 5\u20138 weeks (reduced from 6\u201310 \u2014 admin API already built)', italic=True)

add_heading('Tier 3 \u2014 Enterprise Training Platform ($1,000+/seat)', level=2)

add_bullet('Everything in Tier 2')
add_bullet('Facilitator guides and group discussion materials')
add_bullet('Custom branding per organization')
add_bullet('Industry-specific example packs')
add_bullet('LMS integration (SCORM/xAPI)')
add_bullet('SOC 2 compliance preparation')
add_bullet('Advanced analytics and reporting dashboards')
add_para('Estimated effort: 3\u20136 months', italic=True)

doc.add_page_break()


# ============================================================
# UPDATED SCORECARD
# ============================================================
add_heading('Overall Scorecard', level=1)

add_table(
    ['Dimension', 'v1 Rating', 'v2 Rating', 'Change', 'Notes'],
    [
        ['Curriculum quality', '\u2605\u2605\u2605\u2605\u2606', '\u2605\u2605\u2605\u2605\u2606', '\u2014', 'Unchanged \u2014 well-written, practical, non-generic'],
        ['Pedagogical design', '\u2605\u2605\u2605\u2605\u2606', '\u2605\u2605\u2605\u2605\u2606', '\u2014', 'Unchanged \u2014 strong learn-practice-track cycle'],
        ['Content realism', '\u2605\u2605\u2605\u2605\u2606', '\u2605\u2605\u2605\u2605\u2606', '\u2014', 'Unchanged \u2014 real workplace scenarios across domains'],
        ['Assessment rigor', '\u2605\u2605\u2605\u2606\u2606', '\u2605\u2605\u2605\u2606\u2606', '\u25b2 slight', '4 criteria/lesson now, but still existence-based'],
        ['Enterprise readiness', '\u2605\u2605\u2606\u2606\u2606', '\u2605\u2605\u2606\u2606\u2606', '\u2014', 'No SSO, no cohort UI, no certification'],
        ['Technical quality', '\u2605\u2605\u2605\u2605\u2606', '\u2605\u2605\u2605\u2605\u2606', '\u25b2 slight', 'Schema validation added; exception handling improved'],
        ['Security posture', '\u2605\u2605\u2605\u2606\u2606', '\u2605\u2605\u2605\u2606\u2606', '\u25b2', 'Upload validation and schema constraints added; localStorage remains'],
        ['UX/Design', '\u2605\u2605\u2605\u2605\u2606', '\u2605\u2605\u2605\u2605\u2606', '\u25b2', 'Onboarding overlay and recommended lesson path added'],
        ['Mobile experience', '\u2605\u2605\u2605\u2605\u2606', '\u2605\u2605\u2605\u2605\u2606', '\u2014', 'Unchanged \u2014 thoughtful responsive design'],
        ['Accessibility', '\u2605\u2605\u2605\u2606\u2606', '\u2605\u2605\u2605\u2606\u2606', '\u25b2', 'Focus trapping, lang attr added; confirm() dialogs remain'],
        ['Test coverage', '\u2605\u2605\u2606\u2606\u2606', '\u2605\u2605\u2606\u2606\u2606', '\u25b2 slight', '18 test files, all 12 lessons covered; ~30\u201340% estimated'],
        ['Commercial readiness', '\u2605\u2605\u2606\u2606\u2606', '\u2605\u2605\u2606\u2606\u2606', '\u25b2 slight', 'Onboarding done; payments and certs still missing'],
    ]
)

add_para(
    'No dimension dropped. Five dimensions showed incremental improvement. None crossed a full '
    'star threshold \u2014 the hardening sprint fixed real issues but the commercial blockers '
    '(payments, SSO, certification) remain untouched.',
    italic=True
)

doc.add_page_break()


# ============================================================
# CORRECTIONS TO V1 REVIEW
# ============================================================
add_heading('Corrections to the v1 Review', level=1)

add_para(
    'The independent audit identified several inaccuracies in the original review document. '
    'These are corrected here for the record.'
)

add_table(
    ['v1 Claim', 'Actual Finding', 'Severity'],
    [
        ['"Lesson 11 feels thinner than the others"',
         'Lesson 11 has a 682-line backend (3rd largest of all lessons), extensive Learn tab content, '
         'full CRUD with statistics endpoint, and comparable scope to Lessons 3 and 5. This claim is inaccurate.',
         'Moderate \u2014 unfairly undersells a solid lesson'],
        ['"400+ CSS custom properties"',
         'Approximately 70 unique CSS variables, ~210 total declarations across 3 theme blocks (light, dark, high-contrast). '
         'Well short of 400+.',
         'Minor \u2014 factual overcount'],
        ['"Some inputs use placeholder as label"',
         'All form inputs across the application use proper <label> elements with htmlFor attributes. '
         'Placeholders are used as supplementary hints only.',
         'Minor \u2014 issue does not exist'],
        ['"No lang attribute on HTML"',
         'index.html has lang="en" on the <html> element.',
         'Minor \u2014 was already fixed or never missing'],
        ['"No organizational features"',
         'Admin backend has full cohort management API (create, update, member management, basic stats) '
         'plus A/B experiment infrastructure. The gap is no admin UI, not no backend.',
         'Moderate \u2014 understates existing infrastructure'],
    ]
)

doc.add_page_break()


# ============================================================
# BOTTOM LINE
# ============================================================
add_heading('Bottom Line', level=1)

add_para(
    'The platform-hardening sprint closed real gaps: input validation across all schemas, '
    'accessibility improvements on the mobile sidebar, proper exception handling in auth, '
    'and a genuine onboarding flow. These are meaningful improvements, not cosmetic.',
    space_after=12
)

add_para(
    'However, the commercial blockers identified in v1 remain. No payment system. No certification. '
    'No SSO. Assessment thresholds are still trivially low. These are the same gaps that separate '
    '"good side project" from "sellable product."',
    space_after=12
)

add_para(
    'The shortest path to revenue is still Tier 1 (self-serve individual at $20\u201350/user), '
    'now estimated at 2\u20133 weeks instead of 2\u20134. The core educational product remains '
    'the strongest asset \u2014 genuinely better than most AI training available today. '
    'The curriculum teaches real, practical skills that non-technical knowledge workers actually need.',
    space_after=12
)

add_para(
    'The strategic question is unchanged: sell it as a $30 self-paced course now, or invest in '
    'the enterprise features that unlock $500/seat pricing. The hardening sprint moved the needle '
    'on quality and security, but the revenue infrastructure still needs to be built.',
    space_after=12
)

add_separator()

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('Generated from comprehensive codebase audit \u2014 February 2026 \u2014 Revision 2')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.italic = True


# ============================================================
# SAVE
# ============================================================
output_path = '/root/ClaudeProjects/ai-manager-skills/AI_Manager_Skills_Commercial_Review_v2.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
