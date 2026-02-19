from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# Helper functions
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return h

def add_para(text, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_rich_para(parts, space_after=6):
    """parts is list of (text, bold, italic) tuples"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    for text, bold, italic in parts:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
    return p

def add_bullet(text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()  # spacing after table
    return table

def add_separator():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('─' * 60)
    run.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
    run.font.size = Pt(8)


# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('AI Manager Skills Platform')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Commercial Viability Review')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run(datetime.date.today().strftime('%B %d, %Y'))
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_paragraph()
doc.add_paragraph()

conf = doc.add_paragraph()
conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = conf.add_run('Comprehensive review of curriculum, technology, UX, and commercial readiness')
run.italic = True
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_page_break()


# ============================================================
# EXECUTIVE SUMMARY
# ============================================================
add_heading('Executive Summary', level=1)

add_rich_para([
    ('Could this be sold as an educational tool? ', False, False),
    ('Yes — but it needs work to get there.', True, False),
], space_after=8)

add_para(
    'The foundation is strong. The six managerial concepts are well-defined and teachable. '
    'The interactive tools provide genuine hands-on practice with real AI integration. '
    'The tech stack is production-quality. The gaps are in enterprise readiness, assessment rigor, '
    'and monetization infrastructure — all fixable.'
)

add_para(
    'This review covers four dimensions: curriculum and pedagogy, frontend user experience, '
    'backend architecture and security, and commercial positioning. Each section includes '
    'specific findings, not generalities.'
)

doc.add_page_break()


# ============================================================
# SECTION 1: WHAT'S GENUINELY GOOD
# ============================================================
add_heading("What's Genuinely Good", level=1)

# --- Curriculum ---
add_heading('The Curriculum Is the Real Asset', level=2)

add_para(
    'The six managerial concepts (Context Assembly, Quality Judgment, Task Decomposition, '
    'Iterative Refinement, Workflow Integration, Frontier Recognition) are well-defined, '
    'teachable, and address real pain points. A corporate training buyer would find these '
    'credible — they are not buzzwords.'
)

add_para(
    'Every lesson follows a tight learn-practice-track cycle. Users build actual deliverables '
    '(templates, trust matrices, checklists, delegation plans, reference cards), not just answer '
    'quizzes. The Learn tabs use concrete workplace scenarios across marketing, HR, finance, IT, '
    'and operations — not just tech examples.'
)

add_table(
    ['Module', 'Lessons', 'Primary Concepts'],
    [
        ['Foundation', '1–3', 'Context Assembly'],
        ['Documentation & Trust', '4–6', 'Context Assembly, Quality Judgment'],
        ['Workflow', '7–10', 'Task Decomposition, Iterative Refinement, Workflow Integration'],
        ['Advanced', '11–12', 'Frontier Recognition, Integration'],
    ]
)

# --- Pedagogy ---
add_heading('Pedagogical Design Is Solid', level=2)

add_bullet('Problem/Skill framing: ', 'Immediate relevance — ', 0)
add_para(
    'Each lesson opens with "The Problem" and "The Skill," immediately answering '
    '"why should I care?" before diving into content.'
)

add_bullet('Behavioral assessment: ', 'Measures doing, not knowing — ', 0)
add_para(
    'Self-assessment criteria are behavioral ("analyzed at least 1 conversation," '
    '"identified at least 1 context pattern") rather than knowledge-based ("can you define '
    'context assembly?"). This aligns with adult learning theory.'
)

add_bullet('Cross-lesson integration: ', 'Lessons reference each other — ', 0)
add_para(
    'Lesson 3 pulls data from Lesson 1. Lesson 8 imports from Lesson 7. Lesson 12 aggregates '
    'all lessons. ConnectionCallout components create a narrative thread without forcing '
    'sequential unlocking.'
)

add_bullet('Before/after comparisons: ', 'Show, don\'t tell — ', 0)
add_para(
    'Every lesson includes side-by-side examples of vague vs. specific approaches. '
    'The differences are visceral and memorable.'
)

# --- Tech ---
add_heading('Technology Stack Is Production-Quality', level=2)

add_table(
    ['Layer', 'Technology', 'Notes'],
    [
        ['Backend', 'FastAPI (Python, async)', 'JWT auth, rate limiting, circuit breaker'],
        ['AI Integration', 'Claude API (Anthropic)', 'Real LLM calls with retry logic and graceful degradation'],
        ['Database', 'PostgreSQL + SQLAlchemy 2.0', 'UUID primary keys, per-user data isolation, cascade deletes'],
        ['Frontend', 'React 18 + Vite', 'Custom CSS (no framework), dark/light themes, responsive'],
        ['Deployment', 'Docker Compose', 'Frontend (nginx), backend, postgres, admin panel'],
    ]
)

add_para(
    'This is not a demo or prototype. The AI integration includes a circuit breaker pattern '
    '(opens after 5 failures, half-open after 30s cooldown), exponential backoff retries, '
    'and structured error handling that surfaces user-friendly messages.'
)

doc.add_page_break()


# ============================================================
# SECTION 2: WHAT'S MISSING
# ============================================================
add_heading("What's Missing for a Sellable Product", level=1)

# --- Assessment ---
add_heading('1. Assessment Rigor Is Too Low', level=2)

add_para(
    'Criteria like "created 1 template" are trivially satisfiable. No quality gates exist — '
    'completing a lesson means the artifact exists, not that it is good. There is no competency '
    'proof or certification at the end.'
)

add_bullet('No minimum depth checks (1 template vs. 5 tested templates)')
add_bullet('No rubric for quality of deliverables')
add_bullet('No final integrated assessment beyond the Lesson 12 challenge')
add_bullet('A corporate buyer paying per-seat expects measurable outcomes')

# --- Org features ---
add_heading('2. No Organizational/Team Features', level=2)

add_para(
    'The product is individual-only. There are no cohort dashboards, no peer review, '
    'no team progress visibility, no facilitator guides.'
)

add_bullet('The admin panel exists but is minimal (user management, basic CSV export)')
add_bullet('L&D teams need: cohort comparisons, completion reports, discussion prompts')
add_bullet('No manager view showing "my team\'s progress"')
add_bullet('No group practice or peer learning exercises')

# --- Onboarding ---
add_heading('3. No Onboarding or Guided Path', level=2)

add_para(
    'Users land on the Dashboard and see 12 lessons with no tutorial, no "start here" wizard, '
    'and no "what to do next" recommendations. All lessons are open from day one, which is '
    'flexible but can overwhelm new users.'
)

# --- Security ---
add_heading('4. Security Gaps for Enterprise', level=2)

add_table(
    ['Issue', 'Severity', 'Fix Effort'],
    [
        ['Tokens in localStorage (XSS vulnerable)', 'High', 'Medium — switch to HttpOnly cookies'],
        ['No file size/type validation on uploads', 'Medium', 'Low — add middleware'],
        ['Rate limiting on only 2–3 endpoints', 'Medium', 'Low — extend to all AI endpoints'],
        ['No SSO/SAML integration', 'High (enterprise)', 'High — requires auth refactor'],
        ['No password reset flow', 'Medium', 'Low — add email flow'],
        ['get_optional_user swallows all exceptions', 'Low', 'Low — log unexpected errors'],
    ]
)

# --- Content ---
add_heading('5. Content Gaps', level=2)

add_bullet('Lesson 11 (Frontier Recognition) feels thinner than the others')
add_bullet('No "apply this to your real work this week" structured exercises')
add_bullet('No advanced/stretch path for fast learners')
add_bullet('No industry customization (everyone sees same examples regardless of role)')
add_bullet('No frontier case library from real encounters')

# --- Monetization ---
add_heading('6. No Monetization Infrastructure', level=2)

add_bullet('No payment/subscription system')
add_bullet('No license management or seat counting')
add_bullet('No usage metering')
add_bullet('Guest login creates anonymous accounts with no conversion funnel')

doc.add_page_break()


# ============================================================
# SECTION 3: DETAILED FINDINGS
# ============================================================
add_heading('Detailed Technical Findings', level=1)

# --- Backend ---
add_heading('Backend Architecture', level=2)

add_para(
    'The backend is well-structured with a modular design: each lesson has its own directory '
    'with schemas, routes, and analyzers. All endpoints have real business logic — none are stubs.'
)

add_rich_para([
    ('AI Integration: ', True, False),
    ('Real Claude API calls (claude-3-haiku by default, configurable). Integrated in Lessons 1, 3, 5, and 12. '
     'Circuit breaker tracks failures in 60-second windows, opens after 5 failures, allows test request '
     'after 30s cooldown. Retries up to 3 times with exponential backoff (1s, 2s, 4s). Auth errors never retried.', False, False),
])

add_rich_para([
    ('Authentication: ', True, False),
    ('Production-ready JWT with access tokens (15-min) and refresh tokens (7-day). Passwords hashed with bcrypt. '
     'Rate limiting on register/login (5/minute). Token type validation prevents refresh-as-access attacks. '
     'Refresh token revocation on logout.', False, False),
])

add_rich_para([
    ('Database: ', True, False),
    ('PostgreSQL with SQLAlchemy 2.0 async. UUID primary keys. Per-user data isolation via user_id foreign keys '
     'with CASCADE deletes. 20+ models covering all 12 lessons plus analytics, admin cohorts, and A/B experiments.', False, False),
])

add_rich_para([
    ('Test Coverage: ', True, False),
    ('Approximately 30–40% of codebase. Tests exist for auth flows, Lesson 1, Lesson 3, and Lesson 8. '
     'Uses pytest + pytest-asyncio. Needs significant expansion for production confidence.', False, False),
])

# --- Frontend ---
add_heading('Frontend User Experience', level=2)

add_para(
    'React 18 with modern hooks, React Router v6, Vite build tool. Zero external UI libraries — '
    'all custom CSS (4,000+ lines, well-organized with CSS custom properties for theming).'
)

add_rich_para([
    ('Themes: ', True, False),
    ('Light and dark themes via CSS custom properties (~70 unique variables, ~210 total declarations across '
     '3 theme blocks). High-contrast theme defined in CSS but not yet enabled in code.', False, False),
])

add_rich_para([
    ('Responsive Design: ', True, False),
    ('Three breakpoints (768px, 1024px, 480px). Mobile hamburger menu with slide-in sidebar. '
     'Grid layouts reflow naturally. Touch targets meet WCAG minimum (44px). No horizontal scroll issues.', False, False),
])

add_rich_para([
    ('Accessibility: ', True, False),
    ('Skip link, semantic HTML, aria-expanded on toggles, aria-label on icon buttons, '
     'focus-visible states, reduced-motion support. Gaps: no lang attribute on html, no focus trapping '
     'on mobile sidebar, confirm() dialogs instead of accessible modals, some inputs use placeholder as label.', False, False),
])

# --- Lesson UX ---
add_heading('Lesson Experience (What Users Actually Do)', level=2)

add_table(
    ['Lesson', 'Title', 'What Users Build', 'AI-Powered?'],
    [
        ['1', 'Context Tracker', 'Analyzed conversations with coaching feedback', 'Yes — full analysis'],
        ['2', 'Feedback Analyzer', 'Scored feedback entries with rewrite suggestions', 'Rule-based'],
        ['3', 'Template Builder', 'Reusable prompt templates with test results', 'Yes — template testing'],
        ['4', 'Context Docs', 'Living project documents with AI session prompts', 'Prompt generation'],
        ['5', 'Trust Matrix', 'Personal trust ratings with prediction tracking', 'Yes — calibration insights'],
        ['6', 'Verification Tools', 'Reusable verification checklists', 'CRUD only'],
        ['7', 'Task Decomposer', 'Project breakdowns with AI/human categorization', 'CRUD + analytics'],
        ['8', 'Delegation Tracker', 'Structured delegations with task workflows', 'CRUD + templates'],
        ['9', 'Iteration Passes', 'Multi-pass refinement records with feedback quality', 'Feedback scoring'],
        ['10', 'Status Reporter', 'Workflow templates with execution tracking', 'Prompt generation'],
        ['11', 'Frontier Mapper', 'AI reliability zone maps with encounter logs', 'CRUD only'],
        ['12', 'Reference Card', 'Personal quick-reference aggregating all lessons', 'Yes — card generation'],
    ]
)

doc.add_page_break()


# ============================================================
# SECTION 4: COMPETITIVE POSITIONING
# ============================================================
add_heading('Competitive Positioning', level=1)

add_table(
    ['Factor', 'This Project', 'Typical Competitors'],
    [
        ['Content depth', 'Strong — 12 hands-on lessons', 'Often shallow webinars or slide decks'],
        ['Interactivity', 'Real tools with AI integration', 'Usually just videos + quizzes'],
        ['Personalization', 'Builds personal artifacts', 'Generic one-size-fits-all'],
        ['Enterprise readiness', 'Weak — no SSO, no cohort mgmt', 'Strong in established LMS platforms'],
        ['Price point fit', '$20–50 self-serve, $200–500/seat enterprise', 'Varies widely'],
    ]
)

add_rich_para([
    ('The differentiator is real. ', True, False),
    ('Most "AI training" is either deeply technical (prompt engineering for developers) or superficially '
     'motivational ("AI will change everything!"). This occupies a genuine gap: ', False, False),
    ('practical skill-building for non-technical knowledge workers.', True, False),
    (' That is a real market.', False, False),
])

# --- Who would buy ---
add_heading('Target Buyer Fit', level=2)

add_rich_para([('Strong fit:', True, False)])
add_bullet('Knowledge workers wanting to skill up individually on AI collaboration')
add_bullet('Managers wanting personal AI skills before rolling out to teams')
add_bullet('Career development for non-technical roles moving into AI-adjacent work')
add_bullet('Self-paced learner populations (remote-first, asynchronous)')

add_rich_para([('Weaker fit:', True, False)])
add_bullet('Organizations needing to upskill entire teams at once (no cohort tools)')
add_bullet('Regulated industries needing documented competency verification')
add_bullet('Companies wanting trainer-led cohorts with peer learning')
add_bullet('Groups wanting industry-specific content only')

doc.add_page_break()


# ============================================================
# SECTION 5: WHAT IT WOULD TAKE TO SELL
# ============================================================
add_heading('What It Would Take to Sell', level=1)

# --- Tier 1 ---
add_heading('Tier 1 — Self-Serve Individual Product ($20–50/user)', level=2)

add_bullet('Add Stripe integration and paywall')
add_bullet('Add proper onboarding flow (guided first-lesson experience)')
add_bullet('Fix security basics (HttpOnly cookies, file upload validation)')
add_bullet('Add password reset via email')
add_bullet('Polish guest-to-paid conversion funnel')
add_para('Estimated effort: 2–4 weeks', italic=True)

# --- Tier 2 ---
add_heading('Tier 2 — Team/SMB Product ($200–500/seat)', level=2)

add_bullet('Everything in Tier 1')
add_bullet('Cohort management dashboard for managers/admins')
add_bullet('Completion certificates (PDF generation)')
add_bullet('Raise assessment thresholds (create 3+ items per lesson, quality rubrics)')
add_bullet('Add "apply to your work" structured exercises')
add_bullet('SSO/SAML for enterprise authentication')
add_para('Estimated effort: 6–10 weeks', italic=True)

# --- Tier 3 ---
add_heading('Tier 3 — Enterprise Training Platform ($1,000+/seat)', level=2)

add_bullet('Everything in Tier 2')
add_bullet('Facilitator guides and group discussion materials')
add_bullet('Custom branding per organization')
add_bullet('Industry-specific example packs')
add_bullet('LMS integration (SCORM/xAPI)')
add_bullet('SOC 2 compliance preparation')
add_bullet('Advanced analytics and reporting dashboards')
add_para('Estimated effort: 3–6 months', italic=True)

doc.add_page_break()


# ============================================================
# SECTION 6: SCORECARD
# ============================================================
add_heading('Overall Scorecard', level=1)

add_table(
    ['Dimension', 'Rating', 'Notes'],
    [
        ['Curriculum quality', '★★★★☆', 'Well-written, practical, non-generic. Real differentiator.'],
        ['Pedagogical design', '★★★★☆', 'Scenario-based, hands-on, self-paced. Strong learn-practice-track cycle.'],
        ['Content realism', '★★★★☆', 'Examples match real work situations across multiple domains.'],
        ['Assessment rigor', '★★★☆☆', 'Behavioral criteria are good; thresholds are too low.'],
        ['Enterprise readiness', '★★☆☆☆', 'No SSO, no cohort management, no trainer tools.'],
        ['Technical quality', '★★★★☆', 'Production-ready backend, clean frontend, real AI integration.'],
        ['Security posture', '★★★☆☆', 'Solid foundations but gaps in upload validation and token storage.'],
        ['UX/Design', '★★★★☆', 'Clean, professional, responsive. Needs onboarding.'],
        ['Mobile experience', '★★★★☆', 'Thoughtful responsive design, touch-friendly.'],
        ['Accessibility', '★★★☆☆', 'Good basics (skip links, focus states). Missing modals and focus trapping.'],
        ['Test coverage', '★★☆☆☆', '30–40% coverage. Auth tested; most lessons untested.'],
        ['Commercial readiness', '★★☆☆☆', 'No payments, no onboarding, no conversion funnel.'],
    ]
)

doc.add_page_break()


# ============================================================
# SECTION 7: BOTTOM LINE
# ============================================================
add_heading('Bottom Line', level=1)

add_para(
    'The curriculum and interactive tools are the strongest parts — they are genuinely better than '
    'most AI training available today. The tech stack is solid and production-ready for a v1. The '
    'gaps are all in the business and enterprise layer, not in the core educational product.',
    space_after=12
)

add_para(
    'The shortest path to revenue is Tier 1 (self-serve individual). You have a real product that '
    'teaches real skills. The question is whether you want to sell it as a $30 self-paced course '
    'or invest in the enterprise features that unlock $500/seat pricing.',
    space_after=12
)

add_para(
    'The differentiator — practical AI skill-building for non-technical knowledge workers — is '
    'genuine and underserved. Most competitors offer either developer-focused prompt engineering '
    'or shallow motivational content. This sits in the productive middle ground where the real '
    'market demand lives.',
    space_after=12
)

add_separator()

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('Generated from comprehensive codebase review — February 2026')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.italic = True


# ============================================================
# SAVE
# ============================================================
output_path = '/root/ClaudeProjects/ai-manager-skills/AI_Manager_Skills_Commercial_Review.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
